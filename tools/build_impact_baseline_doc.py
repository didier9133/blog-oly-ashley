from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "seo-impact-baseline-2026-07-13"
OUT_FILE = OUT_DIR / "ashley-leon-seo-analytics-baseline-deconstructing-christianity-en.docx"


BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(31, 77, 120)
INK = RGBColor(35, 42, 51)
GRAY = RGBColor(93, 102, 113)
MUTED = RGBColor(116, 124, 135)
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF5EF"
PALE_AMBER = "FFF5E6"
WHITE = "FFFFFF"
BORDER = "D7DDE5"
GREEN = RGBColor(38, 115, 78)
AMBER = RGBColor(163, 101, 24)
BLUE_HEX = "2E74B5"
NAVY_HEX = "1F4D78"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_repeat_header_text(cell, text: str, color=WHITE) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(8.4)
    r.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_page_field_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("Ashley Leon  |  Baseline BSL-DC-2026-07-13  |  Page ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    page_run = p.add_run()
    add_field(page_run, "PAGE")
    page_run.font.size = Pt(8.5)
    page_run.font.color.rgb = MUTED
    of_run = p.add_run(" of ")
    of_run.font.size = Pt(8.5)
    of_run.font.color.rgb = MUTED
    total_run = p.add_run()
    add_field(total_run, "NUMPAGES")
    total_run.font.size = Pt(8.5)
    total_run.font.color.rgb = MUTED


def set_header(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [4680, 4680]
    for i, width in enumerate(widths):
        set_cell_width(table.cell(0, i), width)
        set_cell_margins(table.cell(0, i), top=0, start=0, bottom=35, end=0)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    rl = left.add_run("ORGANIC SEARCH & ANALYTICS BASELINE")
    rl.bold = True
    rl.font.size = Pt(8)
    rl.font.color.rgb = NAVY
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    rr = right.add_run("DECONSTRUCTING CHRISTIANITY")
    rr.font.size = Pt(8)
    rr.font.color.rgb = MUTED
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": BORDER})


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 10),
        ("Subtitle", 14, GRAY, 0, 16),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    set_header(section)
    set_page_field_footer(section)


def add_text(doc: Document, text: str, *, bold=False, italic=False, color=None, size=None,
             align=None, before=0, after=6, keep=False) -> object:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if keep:
        p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc: Document, text: str, *, level=0, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_number(doc: Document, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    marker = p.add_run(f"{number}.\t")
    marker.bold = True
    p.add_run(text)


def add_callout(doc: Document, title: str, body: str, *, fill=PALE_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(0.1)
    table.columns[1].width = Inches(6.4)
    set_cell_width(table.cell(0, 0), 120)
    set_cell_width(table.cell(0, 1), 9240)
    set_cell_shading(table.cell(0, 0), str(accent))
    set_cell_shading(table.cell(0, 1), fill)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell,
                        top={"val": "single", "sz": "4", "color": BORDER},
                        bottom={"val": "single", "sz": "4", "color": BORDER},
                        left={"val": "single", "sz": "4", "color": BORDER},
                        right={"val": "single", "sz": "4", "color": BORDER})
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = accent
    p2 = table.cell(0, 1).add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    p2.runs[0].font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_metric_strip(doc: Document, metrics: list[tuple[str, str, str]], *, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = 9360 // len(metrics)
    for idx, (value, label, note) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_width(cell, width)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=110, start=130, bottom=100, end=130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell,
                        top={"val": "single", "sz": "4", "color": BORDER},
                        bottom={"val": "single", "sz": "4", "color": BORDER},
                        left={"val": "single", "sz": "4", "color": BORDER},
                        right={"val": "single", "sz": "4", "color": BORDER})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(value)
        r.bold = True
        r.font.size = Pt(19)
        r.font.color.rgb = NAVY
        p2 = cell.add_paragraph(label)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(1)
        p2.runs[0].bold = True
        p2.runs[0].font.size = Pt(8.8)
        p3 = cell.add_paragraph(note)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        p3.runs[0].font.size = Pt(7.6)
        p3.runs[0].font.color.rgb = MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int],
              *, font_size=8.4, header_fill=NAVY_HEX, first_col_bold=False) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    header = table.rows[0]
    repeat_table_header(header)
    prevent_row_split(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        set_repeat_header_text(cell, text)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell,
                        top={"val": "single", "sz": "4", "color": BORDER},
                        bottom={"val": "single", "sz": "4", "color": BORDER},
                        left={"val": "single", "sz": "4", "color": BORDER},
                        right={"val": "single", "sz": "4", "color": BORDER})
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        fill = WHITE if row_idx % 2 == 0 else LIGHT
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_shading(cell, fill)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            r.font.size = Pt(font_size)
            r.font.color.rgb = INK
            if first_col_bold and i == 0:
                r.bold = True
            set_cell_border(cell,
                            top={"val": "single", "sz": "4", "color": BORDER},
                            bottom={"val": "single", "sz": "4", "color": BORDER},
                            left={"val": "single", "sz": "4", "color": BORDER},
                            right={"val": "single", "sz": "4", "color": BORDER})
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section_title(doc: Document, number: str, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(number.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = BLUE
    h = doc.add_paragraph(title, style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    if subtitle:
        add_text(doc, subtitle, color=GRAY, size=10, after=10, keep=True)


def add_source_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Source: ")
    r.bold = True
    r.font.size = Pt(7.8)
    r.font.color.rgb = MUTED
    r2 = p.add_run(text)
    r2.font.size = Pt(7.8)
    r2.font.color.rgb = MUTED


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_document() -> Document:
    doc = Document()
    style_document(doc)

    # Cover
    add_text(doc, "MEASUREMENT BASELINE", bold=True, color=BLUE, size=9.5,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=54, after=18)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Organic Search &\nAnalytics Baseline")
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Deconstructing Christianity Pillar")
    add_text(doc,
             "A pre-launch reference for measuring visibility, engagement, and qualified action after implementation.",
             italic=True, color=GRAY, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=4, after=26)

    divider = doc.add_table(rows=1, cols=3)
    divider.alignment = WD_TABLE_ALIGNMENT.CENTER
    divider.autofit = False
    for idx, width in enumerate([3300, 2760, 3300]):
        set_cell_width(divider.cell(0, idx), width)
        set_cell_margins(divider.cell(0, idx), top=0, bottom=0, start=0, end=0)
    set_cell_border(divider.cell(0, 0), bottom={"val": "single", "sz": "10", "color": BLUE_HEX})
    p = divider.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    rr = p.add_run("BASELINE ID  BSL-DC-2026-07-13")
    rr.bold = True
    rr.font.size = Pt(8)
    rr.font.color.rgb = NAVY
    set_cell_border(divider.cell(0, 2), bottom={"val": "single", "sz": "10", "color": BLUE_HEX})

    add_text(doc, "Prepared for Ashley Leon", bold=True, color=NAVY, size=11.5,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=62, after=3)
    add_text(doc, "Pre-launch cut-off: July 13, 2026", color=GRAY, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_text(doc, "Primary URL: https://www.ashleydianaleon.com/en/deconstructing-christianity",
             color=GRAY, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=48)
    add_callout(
        doc,
        "Purpose of this document",
        "Preserve the starting point before the new pillar is published, so future results can be compared against an explicit, dated reference rather than memory or mixed reporting windows.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_text(doc, "Prepared July 13, 2026  •  Status: pre-launch measurement reference",
             color=MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=0)

    add_page_break(doc)

    # Page 2
    add_section_title(doc, "01", "Executive starting point",
                      "What existed before the new pillar and what the implementation is designed to change.")
    add_callout(
        doc,
        "Baseline conclusion",
        "Before launch, the pillar URL had no production history. Therefore, URL-level rankings, Search Console impressions, GA4 views, and pillar CTA events are recorded as not available (N/A), not as measured zeros.",
        fill=PALE_AMBER,
        accent=AMBER,
    )
    add_metric_strip(doc, [
        ("N/A", "Pillar organic history", "URL not live pre-launch"),
        ("5,000", "Core phrase estimate", "Average monthly searches"),
        ("4", "Pillar CTA events", "New event definitions"),
        ("7 / 28 / 90", "Review cadence", "Days after launch"),
    ])
    doc.add_paragraph("What changes with this implementation", style="Heading 2")
    add_table(
        doc,
        ["Dimension", "Before launch", "New implementation"],
        [
            ["Search destination", "No dedicated English pillar URL", "Dedicated evergreen page at /en/deconstructing-christianity"],
            ["Topic coverage", "Relevant ideas distributed across the site", "One structured guide covering meaning, motivations, lived experience, outcomes, harm, next steps, and FAQs"],
            ["Internal discovery", "No primary navigation path to the pillar", "Navigation, footer/sidebar, workbook, and writing links point to the pillar"],
            ["Search presentation", "No page-specific metadata or canonical", "Dedicated title, description, canonical, Open Graph, and English-only alternates"],
            ["Machine-readable context", "No page-level structured data", "Article, BreadcrumbList, and FAQPage structured data with author/entity signals"],
            ["Action measurement", "No pillar-specific CTA history", "Workbook and community CTA views/clicks measured after analytics consent"],
        ],
        [1900, 3100, 4360],
        font_size=8.2,
        first_col_bold=True,
    )
    doc.add_paragraph("What a successful early result means", style="Heading 2")
    add_bullet(doc, "Technical success: the page is reachable, canonicalized, included in the sitemap, and indexed.", bold_prefix="Technical success:")
    add_bullet(doc, "Search success: non-brand impressions and relevant query coverage begin to appear and trend upward.", bold_prefix="Search success:")
    add_bullet(doc, "Content success: visitors engage with the guide rather than immediately abandoning it.", bold_prefix="Content success:")
    add_bullet(doc, "Business-intent success: qualified readers view and click the workbook or community CTAs.", bold_prefix="Business-intent success:")

    add_page_break(doc)

    # Page 3
    add_section_title(doc, "02", "Existing site traffic baseline",
                      "Two systems are preserved separately because their collection methods and reporting windows differ.")
    doc.add_paragraph("Vercel Analytics — production snapshot", style="Heading 2")
    add_metric_strip(doc, [
        ("95", "Visitors", "Last 7 Days"),
        ("243", "Page views", "Last 7 Days"),
        ("73%", "Bounce rate", "Last 7 Days"),
    ], fill=PALE_GREEN)
    add_table(
        doc,
        ["Top path", "Visitors"],
        [
            ["/en", "65"],
            ["/es", "13"],
            ["/en/workbooks/rebuilding-reverence", "9"],
            ["/", "6"],
        ],
        [7600, 1760],
        font_size=8.8,
        first_col_bold=True,
    )
    add_source_note(doc, "Vercel Analytics production dashboard captured July 11, 2026; filter shown as “Last 7 Days” (chart labels July 4–11).")
    doc.add_paragraph("GA4 — early implementation snapshot", style="Heading 2")
    add_metric_strip(doc, [
        ("4", "Active users", "Jun 14–Jul 11"),
        ("4", "New users", "Jun 14–Jul 11"),
        ("17 sec", "Avg. engagement", "Per active user"),
        ("32", "Events", "All event types"),
    ])
    add_callout(
        doc,
        "Instrumentation proof, not a performance result",
        "GA4 Realtime recorded one blog_product_cta_view and one blog_product_cta_click during validation. This confirms the analytics path can receive custom events; it does not yet establish a reliable conversion rate.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_source_note(doc, "GA4 overview captured July 12, 2026 for June 14–July 11, 2026; GA4 Realtime validation captured July 11, 2026.")
    doc.add_paragraph("How to interpret these numbers", style="Heading 2")
    add_bullet(doc, "Do not reconcile Vercel “visitors” directly with GA4 “active users”; they use different collection rules, consent states, and periods.")
    add_bullet(doc, "The GA4 sample is extremely small and may include implementation or QA activity. Use it as an instrumentation baseline, not a demand benchmark.")
    add_bullet(doc, "The existing workbook path received 9 visitors in the Vercel snapshot. This is a useful supporting reference once the pillar begins sending internal traffic to that page.")

    add_page_break(doc)

    # Page 4
    add_section_title(doc, "03", "Organic search demand baseline",
                      "Keyword Planner estimates establish demand before publication; they do not represent Ashley Leon’s current rankings or traffic.")
    add_callout(
        doc,
        "Important volume rule",
        "“deconstruction christianity” and “deconstruction of christianity” were each reported at 5,000 average monthly searches. They must not be added together as 10,000 because their meanings and search results overlap substantially.",
        fill=PALE_AMBER,
        accent=AMBER,
    )
    add_table(
        doc,
        ["Representative keyword", "Avg. monthly", "Intent", "Ads competition", "Strategic role"],
        [
            ["deconstruction christianity", "5,000", "Informational", "Low", "Highest-volume topic leader"],
            ["deconstruction of christianity", "5,000", "Mixed: informational + commercial", "Low", "Borderline phrase; supports pillar and resource discovery"],
            ["faith deconstruction", "500", "Informational", "Low", "Core semantic variant"],
            ["deconstruction of faith", "500", "Informational", "Low", "Core semantic variant"],
            ["religious deconstruction", "500", "Informational", "Low", "Related definition language; +900% YoY in export"],
            ["what is religious deconstruction", "500", "Informational", "Low", "FAQ / definition intent; +900% YoY"],
            ["what is deconstruction in christianity", "500", "Informational", "Low", "Direct question intent"],
            ["what does deconstructing faith mean", "500", "Informational", "Low", "Meaning / FAQ intent"],
            ["religious trauma workbook", "50", "Commercial / transactional", "High", "Independent workbook cluster; index 80"],
        ],
        [2800, 1050, 1700, 1250, 2560],
        font_size=7.8,
        first_col_bold=True,
    )
    add_source_note(doc, "Google Keyword Planner export dated July 12, 2026; data period July 1, 2025–June 30, 2026; United States; Planner language set to all languages. Intent and clustering were validated through a US/English, personalization-disabled SERP review.")
    doc.add_paragraph("Search Console starting point", style="Heading 2")
    add_callout(
        doc,
        "Not supplied at the pre-launch cut-off",
        "No Google Search Console export was available for this baseline. At launch, capture domain-level performance for the previous 28 days and create an exact-page filter for the new URL. The page itself has no pre-launch URL history because it was not live.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_text(doc, "The keyword research documents opportunity. Search Console will document whether the new page earns visibility and clicks.", italic=True, color=GRAY, size=9.5, after=0)

    add_page_break(doc)

    # Page 5
    add_section_title(doc, "04", "Measurement model",
                      "Measure the full path from search visibility to a qualified next step, without treating a CTA click as a sale.")
    add_table(
        doc,
        ["Stage", "Primary system", "Metrics", "Question answered"],
        [
            ["1. Discovery", "Google Search Console", "Index status, impressions, query count, average position", "Is Google finding and showing the pillar?"],
            ["2. Acquisition", "Search Console + GA4", "Organic clicks, CTR, landing sessions/users", "Are searchers choosing and reaching the page?"],
            ["3. Consumption", "GA4", "Views, active users, engagement time, engaged sessions", "Are visitors using the guide?"],
            ["4. Qualified action", "GA4 custom events", "CTA views, clicks, and click-through rates", "Does the guide move readers toward a relevant resource?"],
            ["5. Business outcome", "Commerce/community platform", "Purchase, download, signup, or join completion", "Did the downstream action complete?"],
        ],
        [1400, 1900, 2800, 3260],
        font_size=8.1,
        first_col_bold=True,
    )
    doc.add_paragraph("Pillar event dictionary", style="Heading 2")
    add_table(
        doc,
        ["Event", "When it fires", "Destination"],
        [
            ["pillar_workbook_cta_view", "At least 50% of the workbook CTA is visible after analytics consent", "Rebuilding Reverence"],
            ["pillar_workbook_cta_click", "A visitor clicks the workbook CTA after analytics consent", "Rebuilding Reverence"],
            ["pillar_community_cta_view", "At least 50% of the community CTA is visible after analytics consent", "The In-Between"],
            ["pillar_community_cta_click", "A visitor clicks the community CTA after analytics consent", "The In-Between"],
        ],
        [3100, 4100, 2160],
        font_size=8.1,
        first_col_bold=True,
    )
    add_text(doc, "Shared event parameters: page_slug, page_type, cta_destination, cta_placement, and locale.", color=GRAY, size=8.7, after=7)
    doc.add_paragraph("Calculated KPIs", style="Heading 2")
    add_table(
        doc,
        ["KPI", "Formula", "Interpretation"],
        [
            ["Organic CTR", "Search clicks ÷ search impressions", "How often searchers choose Ashley’s result when it appears"],
            ["Workbook CTA CTR", "Workbook CTA clicks ÷ workbook CTA views", "Response among visitors who actually saw that CTA"],
            ["Community CTA CTR", "Community CTA clicks ÷ community CTA views", "Response among visitors who actually saw that CTA"],
            ["CTA click rate per page view", "All pillar CTA clicks ÷ pillar page views", "Overall action density for the page"],
        ],
        [2400, 3100, 3860],
        font_size=8.1,
        first_col_bold=True,
    )

    add_page_break(doc)

    # Page 6
    add_section_title(doc, "05", "Comparison cadence and decision rules",
                      "The first weeks validate discovery and measurement; the 90-day view is the first meaningful trend assessment.")
    add_table(
        doc,
        ["Checkpoint", "Capture", "What to evaluate", "Decision rule"],
        [
            ["Launch day", "Production QA + annotation", "200 status, canonical, structured data, sitemap entry, GA4 event receipt after consent", "Fix technical or measurement defects before judging content"],
            ["Day 7", "GSC URL inspection + GA4 page report", "Discovery/indexing, first impressions, real page views, event availability", "If not discoverable, investigate sitemap, canonical, robots, and internal links"],
            ["Day 28", "Exact-page GSC and GA4 reports", "Queries, impressions, clicks, CTR, engagement, CTA views/clicks", "Identify emerging queries and weak content/CTA transitions; avoid conclusions from tiny samples"],
            ["Day 90", "28-day vs prior 28-day trend plus cumulative view", "Visibility trend, query breadth, click growth, engagement quality, assisted resource traffic", "Keep, expand, re-title, re-link, or refine sections based on observed query and action data"],
        ],
        [1200, 2200, 3200, 2760],
        font_size=7.9,
        first_col_bold=True,
    )
    doc.add_paragraph("Directional success indicators", style="Heading 2")
    add_bullet(doc, "Indexing: the exact URL is indexed and Google selects the intended canonical.")
    add_bullet(doc, "Relevance: the page earns impressions for several terms within the deconstruction cluster, not only Ashley’s name.")
    add_bullet(doc, "Momentum: impressions, clicks, and relevant query count trend upward across comparable windows.")
    add_bullet(doc, "Quality: engaged visits and reading behavior improve as traffic grows; bounce rate is interpreted alongside engagement, not alone.")
    add_bullet(doc, "Action: at least one CTA receives legitimate views and clicks from non-test users, allowing CTR to be calculated.")
    add_bullet(doc, "Assistance: the workbook and/or community destinations show referral traffic from the pillar.")
    add_callout(
        doc,
        "Not a ranking guarantee",
        "The 5,000-search estimates describe market demand, not expected traffic. Performance depends on indexing, competition, content quality, authority, result presentation, and time. These indicators are decision thresholds, not contractual forecasts.",
        fill=PALE_AMBER,
        accent=AMBER,
    )
    doc.add_paragraph("Recommended comparison windows", style="Heading 2")
    add_number(doc, 1, "Use exact calendar windows of equal length (for example, the first 28 complete days after launch versus the following 28 complete days).")
    add_number(doc, 2, "Preserve the launch date as an annotation and record any major promotion, email, or social campaign that could change traffic.")
    add_number(doc, 3, "Segment the pillar URL, organic search channel, locale, and device before drawing conclusions.")

    add_page_break(doc)

    # Page 7
    add_section_title(doc, "06", "Baseline scorecard",
                      "This is the table to complete at launch, Day 28, and Day 90. N/A means the metric did not exist for the new page before publication.")
    add_table(
        doc,
        ["Metric", "Pre-launch baseline", "Launch / first capture", "Day 28", "Day 90"],
        [
            ["Pillar indexed", "N/A — URL not live", "________", "________", "________"],
            ["Pillar search impressions", "N/A — no URL history", "________", "________", "________"],
            ["Pillar search clicks", "N/A — no URL history", "________", "________", "________"],
            ["Pillar organic CTR", "N/A — no URL history", "________", "________", "________"],
            ["Relevant query count", "N/A — no URL history", "________", "________", "________"],
            ["Pillar GA4 views", "N/A — page absent", "________", "________", "________"],
            ["Pillar active users", "N/A — page absent", "________", "________", "________"],
            ["Avg. engagement time", "N/A — page absent", "________", "________", "________"],
            ["Workbook CTA views / clicks", "N/A — events absent", "________", "________", "________"],
            ["Community CTA views / clicks", "N/A — events absent", "________", "________", "________"],
            ["Workbook destination visitors", "9 in Vercel 7-day snapshot", "________", "________", "________"],
            ["Site visitors / page views", "95 / 243 in Vercel 7-day snapshot", "________", "________", "________"],
        ],
        [2600, 2400, 1660, 1350, 1350],
        font_size=7.6,
        first_col_bold=True,
    )
    add_text(doc, "For GSC, use the exact page filter. For GA4, use the pillar page path and exclude known internal/testing traffic where possible.", italic=True, color=GRAY, size=8.5, after=10)
    doc.add_paragraph("Minimum reporting view", style="Heading 2")
    add_table(
        doc,
        ["Report", "Filter / breakdown", "Export fields"],
        [
            ["Search Console Performance", "Exact pillar URL; Web search", "Clicks, impressions, CTR, position; queries; country; device"],
            ["GA4 Landing page", "Page path contains /en/deconstructing-christianity", "Sessions, active users, views, engagement, source/medium"],
            ["GA4 Events", "Four pillar_* events", "Event count, users, page path, CTA destination, CTA placement"],
            ["Destination-page traffic", "Rebuilding Reverence + The In-Between", "Sessions/users referred by the pillar and downstream completions if available"],
        ],
        [2400, 3400, 3560],
        font_size=8.1,
        first_col_bold=True,
    )

    add_page_break(doc)

    # Page 8
    add_section_title(doc, "07", "Data quality, launch actions, and evidence",
                      "The baseline is useful only if the post-launch capture uses consistent definitions and preserves known limitations.")
    doc.add_paragraph("Known limitations", style="Heading 2")
    add_bullet(doc, "GA4 data is consent-dependent. Visitors who do not accept analytics will not produce the custom pillar events described here.")
    add_bullet(doc, "The current GA4 sample is too small for stable percentages and may include QA activity.")
    add_bullet(doc, "Vercel and GA4 counts are not expected to match; use each system consistently within its own time series.")
    add_bullet(doc, "Keyword Planner volumes are rounded estimates and close variants can share demand. They are prioritization inputs, not traffic forecasts.")
    add_bullet(doc, "A CTA click is an intent signal, not proof of a workbook purchase or community join. Downstream completion tracking is a separate measurement layer.")
    add_bullet(doc, "Search Console is not real-time, so capture dates should be logged and comparisons should use complete reporting windows.")
    doc.add_paragraph("Launch measurement checklist", style="Heading 2")
    add_number(doc, 1, "Deploy the page to production and record the exact go-live date and time.")
    add_number(doc, 2, "Verify the URL returns 200, the canonical points to the final HTTPS URL, and the English page is indexable.")
    add_number(doc, 3, "Validate Article, BreadcrumbList, and FAQPage structured data and confirm the page appears in the sitemap.")
    add_number(doc, 4, "Submit or refresh the sitemap in Search Console, inspect the exact URL, and request indexing if appropriate.")
    add_number(doc, 5, "Accept analytics consent in a clean test session; validate all four pillar events and their parameters in GA4 Realtime/DebugView.")
    add_number(doc, 6, "Exclude or label QA activity, then capture the launch/first-data column in the scorecard.")
    add_number(doc, 7, "Schedule the Day 7, Day 28, and Day 90 reviews using equal, complete comparison windows.")
    doc.add_paragraph("Evidence register", style="Heading 2")
    add_table(
        doc,
        ["Evidence", "Captured / period", "Used for"],
        [
            ["Vercel Analytics production dashboard", "Captured Jul 11, 2026; Last 7 Days", "Visitors, page views, bounce rate, leading paths and referrers"],
            ["GA4 overview", "Captured Jul 12, 2026; Jun 14–Jul 11", "Active/new users, engagement time, total events"],
            ["GA4 Realtime validation", "Captured Jul 11, 2026", "Proof that existing custom CTA events reached GA4"],
            ["Google Keyword Planner export", "Exported Jul 12, 2026; Jul 2025–Jun 2026", "Demand, trend, Ads competition, and representative keyword volumes"],
            ["SERP and implementation review", "Completed Jul 12–13, 2026", "Intent and cluster overlap plus URL, metadata, schema, internal links, sitemap, and pillar event definitions"],
        ],
        [3200, 2500, 3660],
        font_size=7.9,
        first_col_bold=True,
    )

    # The normal post-table spacer can spill onto an otherwise blank final page.
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        paragraph = doc.paragraphs[-1]
        paragraph._element.getparent().remove(paragraph._element)
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.core_properties.title = "Organic Search & Analytics Baseline — Deconstructing Christianity"
    doc.core_properties.subject = "Pre-launch measurement reference"
    doc.core_properties.author = "Ashley Leon SEO Implementation"
    doc.core_properties.keywords = "SEO, GA4, Search Console, baseline, deconstructing Christianity"
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()
