from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_impact_baseline_doc as base


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "seo-impact-baseline-2026-07-13"
OUT_FILE = OUT_DIR / "ashley-leon-linea-base-seo-analitica-deconstructing-christianity-es.docx"


def set_header_es(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, width in enumerate([4680, 4680]):
        base.set_cell_width(table.cell(0, i), width)
        base.set_cell_margins(table.cell(0, i), top=0, start=0, bottom=35, end=0)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    rl = left.add_run("LÍNEA BASE DE BÚSQUEDA ORGÁNICA Y ANALÍTICA")
    rl.bold = True
    rl.font.size = Pt(8)
    rl.font.color.rgb = base.NAVY
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    rr = right.add_run("DECONSTRUCCIÓN DEL CRISTIANISMO")
    rr.font.size = Pt(8)
    rr.font.color.rgb = base.MUTED
    for cell in table.rows[0].cells:
        base.set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": base.BORDER})


def set_footer_es(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("Ashley Leon  |  Línea base BSL-DC-2026-07-13  |  Página ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = base.MUTED
    page_run = p.add_run()
    base.add_field(page_run, "PAGE")
    page_run.font.size = Pt(8.5)
    page_run.font.color.rgb = base.MUTED
    of_run = p.add_run(" de ")
    of_run.font.size = Pt(8.5)
    of_run.font.color.rgb = base.MUTED
    total_run = p.add_run()
    base.add_field(total_run, "NUMPAGES")
    total_run.font.size = Pt(8.5)
    total_run.font.color.rgb = base.MUTED


def style_document_es(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = base.INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Title", 30, base.NAVY, 0, 10),
        ("Subtitle", 14, base.GRAY, 0, 16),
        ("Heading 1", 16, base.BLUE, 16, 8),
        ("Heading 2", 13, base.BLUE, 12, 6),
        ("Heading 3", 11.5, base.NAVY, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    set_header_es(section)
    set_footer_es(section)


def add_fuente(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Fuente: ")
    r.bold = True
    r.font.size = Pt(7.8)
    r.font.color.rgb = base.MUTED
    r2 = p.add_run(text)
    r2.font.size = Pt(7.8)
    r2.font.color.rgb = base.MUTED


def build_document_es() -> Document:
    doc = Document()
    style_document_es(doc)

    # Portada
    base.add_text(doc, "BASE DE MEDICIÓN", bold=True, color=base.BLUE, size=9.5,
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=54, after=18)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Línea base de búsqueda\norgánica y analítica")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Pilar sobre la deconstrucción del cristianismo")
    base.add_text(
        doc,
        "Una referencia previa al lanzamiento para medir visibilidad, interacción y acciones calificadas después de la implementación.",
        italic=True,
        color=base.GRAY,
        size=11.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=4,
        after=26,
    )

    divider = doc.add_table(rows=1, cols=3)
    divider.alignment = WD_TABLE_ALIGNMENT.CENTER
    divider.autofit = False
    for idx, width in enumerate([3300, 2760, 3300]):
        base.set_cell_width(divider.cell(0, idx), width)
        base.set_cell_margins(divider.cell(0, idx), top=0, bottom=0, start=0, end=0)
    base.set_cell_border(divider.cell(0, 0), bottom={"val": "single", "sz": "10", "color": base.BLUE_HEX})
    p = divider.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ID DE LÍNEA BASE  BSL-DC-2026-07-13")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = base.NAVY
    base.set_cell_border(divider.cell(0, 2), bottom={"val": "single", "sz": "10", "color": base.BLUE_HEX})

    base.add_text(doc, "Preparado para Ashley Leon", bold=True, color=base.NAVY, size=11.5,
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=62, after=3)
    base.add_text(doc, "Fecha de corte previa al lanzamiento: 13 de julio de 2026",
                  color=base.GRAY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    base.add_text(doc, "URL principal: https://www.ashleydianaleon.com/en/deconstructing-christianity",
                  color=base.GRAY, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=48)
    base.add_callout(
        doc,
        "Objetivo de este documento",
        "Conservar el punto de partida antes de publicar el nuevo pilar, para que los resultados futuros se comparen con una referencia explícita y fechada, no con recuerdos ni períodos de reporte mezclados.",
        fill=base.PALE_BLUE,
        accent=base.BLUE,
    )
    base.add_text(doc, "Preparado el 13 de julio de 2026  •  Estado: referencia previa al lanzamiento",
                  color=base.MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=0)

    base.add_page_break(doc)

    # Página 2
    base.add_section_title(
        doc,
        "01",
        "Punto de partida ejecutivo",
        "Qué existía antes del nuevo pilar y qué está diseñada para cambiar la implementación.",
    )
    base.add_callout(
        doc,
        "Conclusión de la línea base",
        "Antes del lanzamiento, la URL del pilar no tenía historial en producción. Por ello, rankings por URL, impresiones de Search Console, vistas de GA4 y eventos de CTA del pilar se registran como no disponibles (N/D), no como ceros medidos.",
        fill=base.PALE_AMBER,
        accent=base.AMBER,
    )
    base.add_metric_strip(doc, [
        ("N/D", "Historial orgánico del pilar", "La URL no estaba publicada"),
        ("5.000", "Estimación de frase central", "Búsquedas mensuales promedio"),
        ("4", "Eventos de CTA del pilar", "Nuevas definiciones"),
        ("7 / 28 / 90", "Cadencia de revisión", "Días después del lanzamiento"),
    ])
    doc.add_paragraph("Qué cambia con esta implementación", style="Heading 2")
    base.add_table(
        doc,
        ["Dimensión", "Antes del lanzamiento", "Nueva implementación"],
        [
            ["Destino de búsqueda", "No había una URL pilar dedicada en inglés", "Página evergreen dedicada en /en/deconstructing-christianity"],
            ["Cobertura temática", "Las ideas relevantes estaban distribuidas por el sitio", "Una guía estructurada sobre significado, motivaciones, experiencia, resultados, daño religioso, próximos pasos y preguntas frecuentes"],
            ["Descubrimiento interno", "No había una ruta principal de navegación hacia el pilar", "Navegación, pie de página/barra lateral, workbook y artículos enlazan al pilar"],
            ["Presentación en buscadores", "No había metadata ni canonical específicos", "Título, descripción, canonical, Open Graph y alternates solo en inglés"],
            ["Contexto legible por máquinas", "No había datos estructurados para esta página", "Datos estructurados Article, BreadcrumbList y FAQPage con señales de autora y entidad"],
            ["Medición de acciones", "No había historial de CTA específico del pilar", "Vistas y clics de CTA del workbook y comunidad medidos después del consentimiento analítico"],
        ],
        [1900, 3100, 4360],
        font_size=8.0,
        first_col_bold=True,
    )
    doc.add_paragraph("Qué significa un buen resultado inicial", style="Heading 2")
    base.add_bullet(doc, "Éxito técnico: la página responde, tiene canonical, figura en el sitemap y está indexada.", bold_prefix="Éxito técnico:")
    base.add_bullet(doc, "Éxito en búsqueda: comienzan a aparecer impresiones no relacionadas con la marca y consultas relevantes, con tendencia ascendente.", bold_prefix="Éxito en búsqueda:")
    base.add_bullet(doc, "Éxito de contenido: las personas interactúan con la guía en lugar de abandonarla inmediatamente.", bold_prefix="Éxito de contenido:")
    base.add_bullet(doc, "Éxito de intención comercial: lectores calificados ven y hacen clic en los CTA del workbook o la comunidad.", bold_prefix="Éxito de intención comercial:")

    base.add_page_break(doc)

    # Página 3
    base.add_section_title(
        doc,
        "02",
        "Línea base del tráfico existente",
        "Los dos sistemas se conservan por separado porque sus métodos de recopilación y períodos de reporte son distintos.",
    )
    doc.add_paragraph("Vercel Analytics — captura de producción", style="Heading 2")
    base.add_metric_strip(doc, [
        ("95", "Visitantes", "Últimos 7 días"),
        ("243", "Vistas de página", "Últimos 7 días"),
        ("73%", "Porcentaje de rebote", "Últimos 7 días"),
    ], fill=base.PALE_GREEN)
    base.add_table(
        doc,
        ["Ruta principal", "Visitantes"],
        [
            ["/en", "65"],
            ["/es", "13"],
            ["/en/workbooks/rebuilding-reverence", "9"],
            ["/", "6"],
        ],
        [7600, 1760],
        font_size=8.8,
        first_col_bold=True,
    )
    add_fuente(doc, "Panel de producción de Vercel Analytics capturado el 11 de julio de 2026; filtro “Last 7 Days” (el gráfico muestra del 4 al 11 de julio).")
    doc.add_paragraph("GA4 — captura inicial de implementación", style="Heading 2")
    base.add_metric_strip(doc, [
        ("4", "Usuarios activos", "14 jun.–11 jul."),
        ("4", "Usuarios nuevos", "14 jun.–11 jul."),
        ("17 s", "Interacción promedio", "Por usuario activo"),
        ("32", "Eventos", "Todos los tipos"),
    ])
    base.add_callout(
        doc,
        "Prueba de instrumentación, no un resultado de rendimiento",
        "GA4 Realtime registró un blog_product_cta_view y un blog_product_cta_click durante la validación. Esto confirma que la ruta analítica puede recibir eventos personalizados; todavía no establece una tasa de conversión confiable.",
        fill=base.PALE_GREEN,
        accent=base.GREEN,
    )
    add_fuente(doc, "Vista general de GA4 capturada el 12 de julio de 2026 para el período 14 de junio–11 de julio de 2026; validación de GA4 Realtime capturada el 11 de julio.")
    doc.add_paragraph("Cómo interpretar estas cifras", style="Heading 2")
    base.add_bullet(doc, "No se deben conciliar directamente los “visitantes” de Vercel con los “usuarios activos” de GA4; usan reglas, estados de consentimiento y períodos distintos.")
    base.add_bullet(doc, "La muestra de GA4 es extremadamente pequeña y puede incluir actividad de implementación o QA. Sirve como línea base de instrumentación, no como referencia de demanda.")
    base.add_bullet(doc, "La ruta del workbook recibió 9 visitantes en la captura de Vercel. Es una referencia útil cuando el pilar comience a enviar tráfico interno hacia esa página.")

    base.add_page_break(doc)

    # Página 4
    base.add_section_title(
        doc,
        "03",
        "Línea base de demanda orgánica",
        "Las estimaciones de Keyword Planner documentan demanda antes de publicar; no representan los rankings ni el tráfico actual de Ashley Leon.",
    )
    base.add_callout(
        doc,
        "Regla importante sobre el volumen",
        "“deconstruction christianity” y “deconstruction of christianity” reportaron 5.000 búsquedas mensuales promedio cada una. No deben sumarse como 10.000 porque sus significados y resultados de búsqueda se solapan considerablemente.",
        fill=base.PALE_AMBER,
        accent=base.AMBER,
    )
    base.add_table(
        doc,
        ["Keyword representativa", "Prom. mensual", "Intención", "Competencia Ads", "Rol estratégico"],
        [
            ["deconstruction christianity", "5.000", "Informativa", "Baja", "Líder temático de mayor volumen"],
            ["deconstruction of christianity", "5.000", "Mixta: informativa + comercial", "Baja", "Frase fronteriza; apoya el pilar y el descubrimiento de recursos"],
            ["faith deconstruction", "500", "Informativa", "Baja", "Variante semántica central"],
            ["deconstruction of faith", "500", "Informativa", "Baja", "Variante semántica central"],
            ["religious deconstruction", "500", "Informativa", "Baja", "Lenguaje relacionado de definición; +900% interanual"],
            ["what is religious deconstruction", "500", "Informativa", "Baja", "Intención de FAQ/definición; +900% interanual"],
            ["what is deconstruction in christianity", "500", "Informativa", "Baja", "Intención de pregunta directa"],
            ["what does deconstructing faith mean", "500", "Informativa", "Baja", "Intención de significado/FAQ"],
            ["religious trauma workbook", "50", "Comercial/transaccional", "Alta", "Clúster independiente del workbook; índice 80"],
        ],
        [2800, 1050, 1700, 1250, 2560],
        font_size=7.6,
        first_col_bold=True,
    )
    add_fuente(doc, "Exportación de Google Keyword Planner del 12 de julio de 2026; período de datos 1 de julio de 2025–30 de junio de 2026; Estados Unidos; idioma del planificador: todos los idiomas. La intención y agrupación se validaron con un laboratorio SERP de Estados Unidos/inglés sin personalización.")
    doc.add_paragraph("Punto de partida en Search Console", style="Heading 2")
    base.add_callout(
        doc,
        "No disponible en la fecha de corte previa al lanzamiento",
        "No se proporcionó una exportación de Google Search Console para esta línea base. En el lanzamiento, se debe capturar el rendimiento del dominio durante los 28 días previos y crear un filtro de página exacta para la nueva URL. La página no tiene historial previo porque aún no estaba publicada.",
        fill=base.PALE_BLUE,
        accent=base.BLUE,
    )
    base.add_text(doc, "La investigación de keywords documenta la oportunidad. Search Console documentará si la nueva página obtiene visibilidad y clics.", italic=True, color=base.GRAY, size=9.5, after=0)

    base.add_page_break(doc)

    # Página 5
    base.add_section_title(
        doc,
        "04",
        "Modelo de medición",
        "Mediremos el recorrido completo desde la visibilidad en búsquedas hasta el siguiente paso calificado, sin tratar un clic de CTA como una venta.",
    )
    base.add_table(
        doc,
        ["Etapa", "Sistema principal", "Métricas", "Pregunta que responde"],
        [
            ["1. Descubrimiento", "Google Search Console", "Estado de indexación, impresiones, número de consultas, posición promedio", "¿Google encuentra y muestra el pilar?"],
            ["2. Adquisición", "Search Console + GA4", "Clics orgánicos, CTR, sesiones/usuarios de entrada", "¿Las personas eligen y visitan la página?"],
            ["3. Consumo", "GA4", "Vistas, usuarios activos, tiempo de interacción, sesiones con interacción", "¿Las personas usan la guía?"],
            ["4. Acción calificada", "Eventos personalizados de GA4", "Vistas de CTA, clics y tasas de clic", "¿La guía conduce hacia un recurso relevante?"],
            ["5. Resultado comercial", "Plataforma de venta/comunidad", "Compra, descarga, registro o incorporación completada", "¿Se completó la acción posterior?"],
        ],
        [1400, 1900, 2800, 3260],
        font_size=7.9,
        first_col_bold=True,
    )
    doc.add_paragraph("Diccionario de eventos del pilar", style="Heading 2")
    base.add_table(
        doc,
        ["Evento", "Cuándo se activa", "Destino"],
        [
            ["pillar_workbook_cta_view", "Cuando al menos 50% del CTA del workbook es visible después del consentimiento analítico", "Rebuilding Reverence"],
            ["pillar_workbook_cta_click", "Cuando una persona hace clic en el CTA del workbook después del consentimiento analítico", "Rebuilding Reverence"],
            ["pillar_community_cta_view", "Cuando al menos 50% del CTA de la comunidad es visible después del consentimiento analítico", "The In-Between"],
            ["pillar_community_cta_click", "Cuando una persona hace clic en el CTA de la comunidad después del consentimiento analítico", "The In-Between"],
        ],
        [3100, 4100, 2160],
        font_size=7.9,
        first_col_bold=True,
    )
    base.add_text(doc, "Parámetros compartidos: page_slug, page_type, cta_destination, cta_placement y locale.", color=base.GRAY, size=8.7, after=7)
    doc.add_paragraph("KPIs calculados", style="Heading 2")
    base.add_table(
        doc,
        ["KPI", "Fórmula", "Interpretación"],
        [
            ["CTR orgánico", "Clics de búsqueda ÷ impresiones de búsqueda", "Con qué frecuencia se elige el resultado de Ashley cuando aparece"],
            ["CTR del CTA del workbook", "Clics del CTA del workbook ÷ vistas del CTA", "Respuesta entre quienes realmente vieron ese CTA"],
            ["CTR del CTA de comunidad", "Clics del CTA de comunidad ÷ vistas del CTA", "Respuesta entre quienes realmente vieron ese CTA"],
            ["Tasa de clic por vista de página", "Todos los clics de CTA ÷ vistas del pilar", "Densidad general de acciones en la página"],
        ],
        [2400, 3100, 3860],
        font_size=7.9,
        first_col_bold=True,
    )

    base.add_page_break(doc)

    # Página 6
    base.add_section_title(
        doc,
        "05",
        "Cadencia de comparación y reglas de decisión",
        "Las primeras semanas validan descubrimiento y medición; la vista a 90 días es la primera evaluación significativa de tendencia.",
    )
    base.add_table(
        doc,
        ["Punto de control", "Captura", "Qué evaluar", "Regla de decisión"],
        [
            ["Día de lanzamiento", "QA de producción + anotación", "Estado 200, canonical, datos estructurados, sitemap y recepción de eventos después del consentimiento", "Corregir defectos técnicos o de medición antes de juzgar el contenido"],
            ["Día 7", "Inspección de URL en GSC + reporte de página en GA4", "Descubrimiento/indexación, primeras impresiones, vistas reales y disponibilidad de eventos", "Si no se descubre, revisar sitemap, canonical, robots y enlaces internos"],
            ["Día 28", "Reportes de página exacta en GSC y GA4", "Consultas, impresiones, clics, CTR, interacción, vistas y clics de CTA", "Detectar consultas emergentes y transiciones débiles; evitar conclusiones con muestras pequeñas"],
            ["Día 90", "Tendencia de 28 días frente a los 28 anteriores + acumulado", "Tendencia de visibilidad, amplitud de consultas, crecimiento de clics, calidad de interacción y tráfico asistido", "Mantener, ampliar, retitular, reenlazar o refinar según consultas y acciones observadas"],
        ],
        [1200, 2200, 3200, 2760],
        font_size=7.6,
        first_col_bold=True,
    )
    doc.add_paragraph("Indicadores direccionales de éxito", style="Heading 2")
    base.add_bullet(doc, "Indexación: la URL exacta está indexada y Google selecciona el canonical previsto.")
    base.add_bullet(doc, "Relevancia: la página obtiene impresiones para varios términos del clúster, no solo para el nombre de Ashley.")
    base.add_bullet(doc, "Impulso: impresiones, clics y cantidad de consultas relevantes crecen en ventanas comparables.")
    base.add_bullet(doc, "Calidad: las visitas con interacción y el comportamiento de lectura mejoran; el rebote se interpreta junto con la interacción.")
    base.add_bullet(doc, "Acción: al menos un CTA recibe vistas y clics legítimos de usuarios no relacionados con pruebas, permitiendo calcular el CTR.")
    base.add_bullet(doc, "Asistencia: los destinos del workbook o la comunidad muestran tráfico referido desde el pilar.")
    base.add_callout(
        doc,
        "No es una garantía de posicionamiento",
        "Las estimaciones de 5.000 búsquedas describen demanda de mercado, no tráfico esperado. El rendimiento depende de indexación, competencia, calidad del contenido, autoridad, presentación del resultado y tiempo. Estos indicadores son criterios de decisión, no pronósticos contractuales.",
        fill=base.PALE_AMBER,
        accent=base.AMBER,
    )
    doc.add_paragraph("Ventanas de comparación recomendadas", style="Heading 2")
    base.add_number(doc, 1, "Usar ventanas de calendario completas y de la misma duración; por ejemplo, los primeros 28 días completos después del lanzamiento frente a los siguientes 28.")
    base.add_number(doc, 2, "Conservar la fecha de lanzamiento como anotación y registrar promociones, correos o campañas sociales que puedan alterar el tráfico.")
    base.add_number(doc, 3, "Segmentar por URL del pilar, canal de búsqueda orgánica, locale y dispositivo antes de sacar conclusiones.")

    # Página 7
    base.add_section_title(
        doc,
        "06",
        "Cuadro de seguimiento de la línea base",
        "Esta tabla se completa en el lanzamiento, el día 28 y el día 90. N/D significa que la métrica no existía para la nueva página antes de publicarla.",
    )
    base.add_table(
        doc,
        ["Métrica", "Línea base previa", "Lanzamiento / 1.ª captura", "Día 28", "Día 90"],
        [
            ["Pilar indexado", "N/D — URL no publicada", "________", "________", "________"],
            ["Impresiones del pilar", "N/D — sin historial", "________", "________", "________"],
            ["Clics orgánicos del pilar", "N/D — sin historial", "________", "________", "________"],
            ["CTR orgánico del pilar", "N/D — sin historial", "________", "________", "________"],
            ["Cantidad de consultas relevantes", "N/D — sin historial", "________", "________", "________"],
            ["Vistas del pilar en GA4", "N/D — página ausente", "________", "________", "________"],
            ["Usuarios activos del pilar", "N/D — página ausente", "________", "________", "________"],
            ["Tiempo promedio de interacción", "N/D — página ausente", "________", "________", "________"],
            ["Vistas / clics CTA workbook", "N/D — eventos ausentes", "________", "________", "________"],
            ["Vistas / clics CTA comunidad", "N/D — eventos ausentes", "________", "________", "________"],
            ["Visitantes al destino del workbook", "9 en captura Vercel de 7 días", "________", "________", "________"],
            ["Visitantes / vistas del sitio", "95 / 243 en captura Vercel de 7 días", "________", "________", "________"],
        ],
        [2600, 2400, 1660, 1350, 1350],
        font_size=7.3,
        first_col_bold=True,
    )
    base.add_text(doc, "En GSC, usar el filtro de página exacta. En GA4, usar la ruta del pilar y excluir el tráfico interno o de pruebas conocido cuando sea posible.", italic=True, color=base.GRAY, size=8.5, after=10)
    doc.add_paragraph("Vista mínima para reportar", style="Heading 2")
    base.add_table(
        doc,
        ["Reporte", "Filtro / desglose", "Campos de exportación"],
        [
            ["Rendimiento de Search Console", "URL exacta del pilar; búsqueda web", "Clics, impresiones, CTR, posición; consultas; país; dispositivo"],
            ["Página de destino en GA4", "Ruta contiene /en/deconstructing-christianity", "Sesiones, usuarios activos, vistas, interacción, fuente/medio"],
            ["Eventos de GA4", "Los cuatro eventos pillar_*", "Número de eventos, usuarios, ruta, destino y ubicación del CTA"],
            ["Tráfico a páginas destino", "Rebuilding Reverence + The In-Between", "Sesiones/usuarios referidos por el pilar y acciones posteriores si están disponibles"],
        ],
        [2400, 3400, 3560],
        font_size=7.9,
        first_col_bold=True,
    )

    base.add_page_break(doc)

    # Página 8
    base.add_section_title(
        doc,
        "07",
        "Calidad de datos, acciones de lanzamiento y evidencia",
        "La línea base solo es útil si la captura posterior usa definiciones consistentes y conserva las limitaciones conocidas.",
    )
    doc.add_paragraph("Limitaciones conocidas", style="Heading 2")
    base.add_bullet(doc, "Los datos de GA4 dependen del consentimiento. Quienes no acepten analítica no producirán los eventos personalizados descritos.")
    base.add_bullet(doc, "La muestra actual de GA4 es demasiado pequeña para porcentajes estables y puede incluir actividad de QA.")
    base.add_bullet(doc, "No se espera que Vercel y GA4 coincidan; cada sistema debe usarse de forma consistente dentro de su propia serie temporal.")
    base.add_bullet(doc, "Los volúmenes de Keyword Planner son estimaciones redondeadas y las variantes cercanas pueden compartir demanda. Sirven para priorizar, no para pronosticar tráfico.")
    base.add_bullet(doc, "Un clic de CTA indica intención, no prueba una compra del workbook ni una incorporación a la comunidad. El seguimiento de la finalización es otra capa de medición.")
    base.add_bullet(doc, "Search Console no funciona en tiempo real; deben registrarse las fechas de captura y usarse ventanas completas.")
    doc.add_paragraph("Lista de verificación para medir el lanzamiento", style="Heading 2")
    base.add_number(doc, 1, "Publicar la página en producción y registrar la fecha y hora exactas de salida.")
    base.add_number(doc, 2, "Verificar respuesta 200, canonical hacia la URL HTTPS final e indexabilidad de la página en inglés.")
    base.add_number(doc, 3, "Validar Article, BreadcrumbList y FAQPage, y confirmar que la página aparece en el sitemap.")
    base.add_number(doc, 4, "Enviar o actualizar el sitemap en Search Console, inspeccionar la URL exacta y solicitar indexación cuando corresponda.")
    base.add_number(doc, 5, "Aceptar analítica en una sesión limpia y validar los cuatro eventos y parámetros en GA4 Realtime/DebugView.")
    base.add_number(doc, 6, "Excluir o etiquetar la actividad de QA y completar la columna de lanzamiento/primera captura del cuadro de seguimiento.")
    base.add_number(doc, 7, "Programar revisiones para los días 7, 28 y 90 con ventanas completas y equivalentes.")
    doc.add_paragraph("Registro de evidencia", style="Heading 2")
    base.add_table(
        doc,
        ["Evidencia", "Captura / período", "Uso"],
        [
            ["Panel de producción de Vercel Analytics", "Capturado 11 jul. 2026; últimos 7 días", "Visitantes, vistas, rebote, rutas y referentes principales"],
            ["Vista general de GA4", "Capturado 12 jul. 2026; 14 jun.–11 jul.", "Usuarios activos/nuevos, interacción y eventos totales"],
            ["Validación de GA4 Realtime", "Capturado 11 jul. 2026", "Prueba de que eventos CTA existentes llegaron a GA4"],
            ["Exportación de Google Keyword Planner", "Exportado 12 jul. 2026; jul. 2025–jun. 2026", "Demanda, tendencia, competencia Ads y volúmenes representativos"],
            ["Revisión SERP y de implementación", "Completada 12–13 jul. 2026", "Intención y clústeres, URL, metadata, schema, enlaces, sitemap y eventos"],
        ],
        [3200, 2500, 3660],
        font_size=7.5,
        first_col_bold=True,
    )

    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        paragraph = doc.paragraphs[-1]
        paragraph._element.getparent().remove(paragraph._element)
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document_es()
    doc.core_properties.title = "Línea base de búsqueda orgánica y analítica — Deconstrucción del cristianismo"
    doc.core_properties.subject = "Referencia de medición previa al lanzamiento"
    doc.core_properties.author = "Implementación SEO de Ashley Leon"
    doc.core_properties.keywords = "SEO, GA4, Search Console, línea base, deconstrucción del cristianismo"
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()
