from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "ga4-eventos-2026-07-17"
OUT_FILE = OUT_DIR / "inventario-eventos-ga4-ashley-leon.docx"

INK = RGBColor(36, 42, 48)
NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(96, 105, 114)
WHITE = RGBColor(255, 255, 255)
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
BORDER = "CBD3DC"
ROW_BORDER = "DDE3EA"


ACTIVE_EVENTS = [
    (
        "blog_product_cta_view",
        "Contenido",
        "Cuando al menos 50% de un CTA de producto dentro de un artículo entra en el área visible.",
        "Registra una exposición relevante al bloque promocional de un producto dentro del blog.",
    ),
    (
        "blog_product_cta_click",
        "Contenido",
        "Cuando la persona hace clic en el enlace principal del CTA de producto de un artículo.",
        "Registra interés activo en pasar del contenido editorial a la página del producto recomendado.",
    ),
    (
        "pillar_workbook_cta_view",
        "Pilar SEO",
        "Cuando al menos 50% del CTA del workbook del pilar entra en el área visible.",
        "Mide cuántas personas llegan a ver la oferta del workbook desde la página pilar.",
    ),
    (
        "pillar_workbook_cta_click",
        "Pilar SEO",
        "Cuando la persona hace clic en el CTA del workbook dentro de la página pilar.",
        "Mide la intención de continuar desde el pilar hacia el recurso Rebuilding Reverence.",
    ),
    (
        "pillar_community_cta_view",
        "Pilar SEO",
        "Cuando al menos 50% del CTA de comunidad del pilar entra en el área visible.",
        "Registra la exposición a la invitación de comunidad ubicada al final de la página pilar.",
    ),
    (
        "pillar_community_cta_click",
        "Pilar SEO",
        "Cuando la persona hace clic en el CTA de comunidad dentro de la página pilar.",
        "Mide el interés en pasar del contenido pilar a la experiencia de comunidad The In-Between.",
    ),
    (
        "newsletter_signup",
        "Newsletter",
        "Después de que la suscripción enviada desde el formulario del pie de página termina correctamente.",
        "Registra una alta confirmada al newsletter y permite atribuirla a la ubicación del formulario.",
    ),
    (
        "view_item",
        "Ecommerce",
        "Al cargar una página de workbook, círculo o comunidad que incorpora el componente de analítica de producto.",
        "Registra la visualización de un producto u oferta para analizar el inicio del embudo comercial.",
    ),
    (
        "begin_checkout",
        "Ecommerce",
        "En el primer intento de enviar el formulario de pago, antes de confirmar el pago con Stripe.",
        "Registra el inicio efectivo del proceso de checkout para un producto determinado.",
    ),
    (
        "purchase",
        "Ecommerce",
        "Cuando se muestra la página de éxito y existe una compra completada asociada a la transacción.",
        "Registra una compra confirmada con su identificador, valor, moneda y producto.",
    ),
]

PARAMETER_GROUPS = [
    (
        "CTA de artículos",
        "article_slug, article_category, primary_keyword (opcional), cta_product, cta_placement y locale.",
    ),
    (
        "CTA del pilar",
        "page_slug, page_type, cta_destination, cta_placement y locale.",
    ),
    (
        "Newsletter",
        "source_location y locale. El correo electrónico no se envía a GA4.",
    ),
    (
        "Ecommerce",
        "locale, value, currency e items (item_id, item_name, item_category); purchase añade transaction_id.",
    ),
]

DECLARED_ONLY = [
    ("related_post_click", "Reservado para clics en contenido relacionado."),
    ("community_cta_click", "Reservado para clics en un CTA general de comunidad."),
    ("circle_cta_click", "Reservado para clics en un CTA del círculo."),
    ("newsletter_cta_click", "Reservado para clics en un CTA que dirige al newsletter."),
]


def set_run_font(run, *, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, *, color=ROW_BORDER, size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = tc_borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], *, indent=120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("INVENTARIO DE MEDICIÓN  |  GA4")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Ashley Leon  |  Página ")
    set_run_font(r, size=8.5, color=MUTED)
    page = p.add_run()
    add_field(page, "PAGE")
    set_run_font(page, size=8.5, color=MUTED)
    r = p.add_run(" de ")
    set_run_font(r, size=8.5, color=MUTED)
    total = p.add_run()
    add_field(total, "NUMPAGES")
    set_run_font(total, size=8.5, color=MUTED)


def add_metadata_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10, color=INK, bold=True)
    value_run = p.add_run(value)
    set_run_font(value_run, size=10, color=INK)


def write_cell(cell, text: str, *, header=False, event_name=False, size=8.7) -> None:
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    if header:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(
        run,
        name="Calibri",
        size=size,
        color=WHITE if header else INK,
        bold=header or event_name,
    )


def add_data_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], *, event_col=False, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    repeat_header(hdr)
    for idx, header in enumerate(headers):
        set_cell_shading(hdr.cells[idx], NAVY.__str__())
        write_cell(hdr.cells[idx], header, header=True, size=8.5)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if row_idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "F8FAFC")
        for idx, value in enumerate(values):
            write_cell(row.cells[idx], value, event_name=event_col and idx == 0, size=font_size)
    set_table_geometry(table, widths)
    return table


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(4)
    kicker.paragraph_format.space_after = Pt(3)
    r = kicker.add_run("REFERENCIA DE ANALÍTICA")
    set_run_font(r, size=9, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("Eventos rastreados en GA4")
    set_run_font(r, size=24, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("Inventario técnico con descripciones provisionales para documentación interna")
    set_run_font(r, size=12.5, color=MUTED, italic=True)

    add_metadata_line(doc, "Proyecto", "Ashley Leon — blog, comunidad y productos digitales")
    add_metadata_line(doc, "Fecha de revisión", "17 de julio de 2026")
    add_metadata_line(doc, "Alcance", "Eventos enviados explícitamente desde la aplicación mediante la integración de GA4")

    doc.add_paragraph("Alcance y criterio", style="Heading 1")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.add_run(
        "Este inventario incluye los eventos con un punto de envío activo en el código. Las descripciones son provisionales. "
        "Todos requieren consentimiento analítico y la etiqueta solo se carga en producción. No se incluyen eventos automáticos "
        "o de Medición mejorada, cuya configuración debe confirmarse directamente en la propiedad de GA4."
    )

    doc.add_paragraph("Eventos activos", style="Heading 1")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run("Se identificaron 10 eventos con disparadores activos en la aplicación.")
    add_data_table(
        doc,
        ["Evento", "Área", "Cuándo se dispara", "Descripción dummy"],
        ACTIVE_EVENTS,
        [2700, 1180, 2340, 3140],
        event_col=True,
        font_size=8.2,
    )

    doc.add_paragraph("Parámetros enviados", style="Heading 1")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(
        "Los valores ausentes se eliminan antes del envío. La instrumentación evita incluir correo electrónico, nombre u otros datos personales en GA4."
    )
    add_data_table(
        doc,
        ["Grupo", "Parámetros principales"],
        PARAMETER_GROUPS,
        [2440, 6920],
        font_size=9.0,
    )

    doc.add_paragraph("Eventos declarados, sin disparador activo encontrado", style="Heading 1")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(
        "Estos nombres están permitidos por la capa de analítica, pero no tienen una llamada de envío localizada. No deben reportarse todavía como activos."
    )
    add_data_table(
        doc,
        ["Evento reservado", "Uso provisional"],
        DECLARED_ONLY,
        [3600, 5760],
        event_col=True,
        font_size=9.0,
    )

    doc.core_properties.title = "Eventos rastreados en GA4"
    doc.core_properties.subject = "Inventario de eventos de analítica con descripciones provisionales"
    doc.core_properties.author = "Ashley Leon"
    doc.core_properties.keywords = "GA4, analítica, eventos, ecommerce, CTA, newsletter"
    return doc


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUT_FILE)
    print(OUT_FILE)
